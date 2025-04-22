import os
import zipfile
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side

logo_path = os.path.join(os.path.dirname(__file__), "logo.png")

# --------- EXCEL EXPORT ---------
def export_to_excel(cv_data, filename):
    wb = Workbook()
    ws = wb.active
    ws.title = "Resume"

    bold = Font(bold=True)
    center = Alignment(horizontal="left", vertical="top", wrap_text=True)
    fill_header = PatternFill(start_color="DCE6F1", end_color="DCE6F1", fill_type="solid")
    border = Border(left=Side(style="thin"), right=Side(style="thin"),
                    top=Side(style="thin"), bottom=Side(style="thin"))

    def write_row(section, content):
        row = [section, content]
        ws.append(row)
        for col in range(1, 3):
            cell = ws.cell(row=ws.max_row, column=col)
            cell.alignment = center
            cell.border = border
            if col == 1:
                cell.font = bold
                cell.fill = fill_header

    pi = cv_data["Personal Information"]
    for k, v in pi.items():
        write_row(k, v)

    for edu in cv_data["Education"]:
        write_row("Education", f"{edu['Degree']} in {edu['Field']} - {edu['Institution']} ({edu['Graduation Year']})")

    write_row("Languages", ", ".join(cv_data["Languages"]))

    for exp in cv_data["Professional Experience"]:
        write_row("Experience", f"{exp['Position']} at {exp['Company']} ({exp['Total Years in Company']} yrs)")
        write_row("Responsibilities", exp["Achievements and Responsibilities"])

    for ach in cv_data["Other Achievements"]:
        write_row("Achievement", f"{ach['Type']}: {ach['Title']} - {ach['Institution']} ({ach['Year']})")

    cc = cv_data["Current Compensation"]
    write_row("Gross Salary", cc["Gross Salary"])
    write_row("Net Salary", cc["Net Salary"])
    write_row("Compensation Type", cc["Compensation Type"])

    ws.column_dimensions["A"].width = 25
    ws.column_dimensions["B"].width = 80

    wb.save(filename)

# --------- WORD EXPORT ---------
def export_to_word(cv_data, filename):
    doc = Document()
    section = doc.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0]

    # Insert logo
    if os.path.exists(logo_path):
        header_paragraph.add_run().add_picture(logo_path, width=Inches(1))
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph("")
    title = doc.add_paragraph("TalentWise HR – Resume Standardized")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(14)

    doc.add_paragraph("")

    def section_title(text):
        p = doc.add_paragraph(text)
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(12)

    pi = cv_data["Personal Information"]
    section_title("👤 Personal Information")
    for k, v in pi.items():
        doc.add_paragraph(f"{k}: {v}")

    section_title("🎓 Education")
    for edu in cv_data["Education"]:
        doc.add_paragraph(f"{edu['Degree']} in {edu['Field']} - {edu['Institution']} ({edu['Graduation Year']})")

    section_title("🌐 Languages")
    doc.add_paragraph(", ".join(cv_data["Languages"]))

    section_title("💼 Professional Experience")
    for exp in cv_data["Professional Experience"]:
        doc.add_paragraph(f"{exp['Position']} at {exp['Company']} ({exp['Total Years in Company']} yrs)")
        doc.add_paragraph(f"Responsibilities: {exp['Achievements and Responsibilities']}")

    section_title("🏅 Other Achievements")
    for ach in cv_data["Other Achievements"]:
        doc.add_paragraph(f"{ach['Type']}: {ach['Title']} - {ach['Institution']} ({ach['Year']})")

    section_title("💰 Current Compensation")
    cc = cv_data["Current Compensation"]
    doc.add_paragraph(f"Gross Salary: {cc['Gross Salary']}")
    doc.add_paragraph(f"Net Salary: {cc['Net Salary']}")
    doc.add_paragraph(f"Compensation Type: {cc['Compensation Type']}")

    doc.save(filename)

# --------- PDF EXPORT ---------
def export_to_pdf(cv_data, filename):
    c = canvas.Canvas(filename, pagesize=letter)
    width, height = letter
    y = height - 60

    if os.path.exists(logo_path):
        c.drawImage(ImageReader(logo_path), 40, y - 20, width=60, preserveAspectRatio=True, mask='auto')
        y -= 10

    c.setFont("Helvetica-Bold", 16)
    c.drawString(110, y, "TalentWise HR")
    c.setFont("Helvetica", 10)
    c.drawString(110, y - 18, "Resume Standardized")
    y -= 45

    def draw_text(label, value, spacing=15):
        nonlocal y
        if y < 100:
            c.showPage()
            y = height - 50
        c.setFont("Helvetica-Bold", 10)
        c.drawString(40, y, f"{label}:")
        c.setFont("Helvetica", 10)
        c.drawString(160, y, str(value))
        y -= spacing

    pi = cv_data["Personal Information"]
    draw_text("👤 Section", "Personal Information", spacing=20)
    for k, v in pi.items():
        draw_text(k, v)

    draw_text("🎓 Section", "Education", spacing=20)
    for edu in cv_data["Education"]:
        draw_text("Education", f"{edu['Degree']} in {edu['Field']} - {edu['Institution']} ({edu['Graduation Year']})")

    draw_text("🌐 Languages", ", ".join(cv_data["Languages"]))

    draw_text("💼 Section", "Professional Experience", spacing=20)
    for exp in cv_data["Professional Experience"]:
        draw_text("Experience", f"{exp['Position']} at {exp['Company']} ({exp['Total Years in Company']} yrs)")
        draw_text("Responsibilities", exp["Achievements and Responsibilities"])

    draw_text("🏅 Section", "Other Achievements", spacing=20)
    for ach in cv_data["Other Achievements"]:
        draw_text("Achievement", f"{ach['Type']}: {ach['Title']} - {ach['Institution']} ({ach['Year']})")

    cc = cv_data["Current Compensation"]
    draw_text("💰 Section", "Current Compensation", spacing=20)
    draw_text("Gross Salary", cc["Gross Salary"])
    draw_text("Net Salary", cc["Net Salary"])
    draw_text("Compensation Type", cc["Compensation Type"])

    c.save()

# --------- ZIP EXPORT ---------
def zip_ready_cvs(cvs_dict, zip_path="exports/ready_cvs.zip"):
    os.makedirs("exports", exist_ok=True)
    with zipfile.ZipFile(zip_path, "w") as zipf:
        for uid, cv in cvs_dict.items():
            if cv["status"] == "ready":
                filename_base = f"{cv['filename'].replace(' ', '_')}_{uid}"

                xlsx = f"exports/{filename_base}.xlsx"
                export_to_excel(cv["data"], xlsx)
                zipf.write(xlsx, os.path.basename(xlsx))

                docx = f"exports/{filename_base}.docx"
                export_to_word(cv["data"], docx)
                zipf.write(docx, os.path.basename(docx))

                pdf = f"exports/{filename_base}.pdf"
                export_to_pdf(cv["data"], pdf)
                zipf.write(pdf, os.path.basename(pdf))
    return zip_path
